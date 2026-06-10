# -*- coding: utf-8 -*-
"""
Генератор звітів (пояснювальних записок) до лабораторних робіт
з дисципліни «Інформаційні технології моніторингу та аналізу даних»

Генерує 6 DOCX файлів за шаблоном ВНТУ (кафедра САІТ)
Використання: python generate_report.py
"""

import os
import sys
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─────────────────────────────────────────────
#  Конфігурація: змініть дані студента тут
# ─────────────────────────────────────────────
STUDENT_FULL_NAME = "Столовник В. О."
STUDENT_FULL_NAME_CAPS = "Володимир СТОЛОВНИК"
STUDENT_GROUP = "2ІСТ-24б"
STUDENT_COURSE = "2"
SUPERVISOR_NAME = "Арсен ЛОСЕНКО"
SUPERVISOR_TITLE = "асистент кафедри САІТ"
SUPERVISOR_DEGREE = "Ph.D."
DEPARTMENT = "САІТ"
DISCIPLINE = "Інформаційні технології моніторингу та аналізу даних"
YEAR = "2025"
CITY = "Вінниця"
UNIVERSITY_SHORT = "ВНТУ"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─────────────────────────────────────────────
#  Визначення лабораторних робіт
# ─────────────────────────────────────────────

LABS = [
    {
        'number': 1,
        'theme': 'EDA – розвідувальний аналіз даних',
        'goal': (
            'ознайомлення з методами розвідувального аналізу даних '
            '(Exploratory Data Analysis, EDA), застосування основних методів EDA '
            'для дослідження та підготовки даних до прогнозування, '
            'набуття навичок аналізу наборів даних, виявлення закономірностей, '
            'взаємозв\'язків між змінними та їх візуалізації'
        ),
        'steps': [
            'Завантажено датасет California Housing із бібліотеки sklearn.datasets.',
            'Проведено первинний аналіз даних: перегляд структури набору даних '
            '(кількість рядків і стовпців, назви змінних, їх типи), '
            'аналіз пропущених значень, визначення розподілу змінних.',
            'Визначено основні статистичні характеристики '
            '(середнє значення, медіана, мінімальне та максимальне значення, '
            'стандартне відхилення).',
            'Побудовано візуалізації: гістограми розподілу числових змінних, '
            'діаграми розсіювання (scatter plots), кореляційну матрицю (heatmap), '
            'box plots для виявлення викидів, pairplot ключових ознак.',
            'Виконано обробку даних: перевірку пропущених значень, '
            'виявлення та обробку викидів методом IQR, '
            'масштабування числових змінних (StandardScaler).',
            'Сформульовано висновки щодо якості та повноти даних.',
        ],
        'conclusion': (
            'Проведено розвідувальний аналіз даних (EDA) для датасету California Housing. '
            'Датасет містить 20640 записів та 9 ознак (8 числових + 1 цільова). '
            'Пропущених значень не виявлено. Кореляційний аналіз показав, що найбільший '
            'вплив на ціну мають: MedInc (медіанний дохід) та AveRooms (середня кількість кімнат). '
            'Виявлено та оброблено викиди. Виконано масштабування всіх числових ознак. '
            'Дані підготовлені для побудови моделей машинного навчання.'
        ),
        'plots_dir': 'lab1_plots',
        'script_name': 'lab1_eda.py',
    },
    {
        'number': 2,
        'theme': 'FE – методи вибору ознак (feature engineering)',
        'goal': (
            'вивчення різних методів вибору ознак (feature selection) '
            'в машинному навчанні, практичне застосування основних методів '
            'відбору ознак у вирішенні завдань прогнозування'
        ),
        'steps': [
            'Завантажено датасет California Housing, виконано розділення на '
            'навчальну та тестову вибірки (80/20).',
            'Візуалізовано важливість ознак (feature importance) за допомогою '
            'алгоритмів Random Forest та Gradient Boosting.',
            'Проаналізовано кореляцію між ознаками та цільовою змінною (Pearson).',
            'Застосовано вбудовані методи вибору ознак (embedded methods): '
            'Lasso-регресія з аналізом коефіцієнтів.',
            'Використано wrapper метод — рекурсивний вибір ознак (RFE) '
            'з LinearRegression як базовим оцінювачем.',
            'Застосовано filter методи: SelectKBest з f_regression та mutual_info_regression.',
            'Проведено порівняння ефективності різних методів вибору ознак '
            'за метриками якості моделі (R², MAE).',
        ],
        'conclusion': (
            'Проведено аналіз методів вибору ознак для датасету California Housing. '
            'Найбільш інформативна ознака — MedInc (медіанний дохід). '
            'Порівняно вбудовані (Lasso), wrapper (RFE) та filter '
            '(SelectKBest з f_regression та Mutual Information) методи. '
            'Для максимальної точності рекомендується використовувати всі ознаки '
            'або Top-5 за RF importance. Для зменшення складності моделі — RFE або Lasso відбір.'
        ),
        'plots_dir': 'lab2_plots',
        'script_name': 'lab2_feature_engineering.py',
    },
    {
        'number': 3,
        'theme': 'Застосування моделей машинного навчання за допомогою бібліотеки scikit-learn',
        'goal': (
            'ознайомлення з основами використання бібліотеки scikit-learn '
            'для реалізації моделей машинного навчання, побудова, тренування '
            'та оцінювання регресійних і класифікаційних моделей'
        ),
        'steps': [
            'Імпортовано бібліотеки scikit-learn, pandas, numpy, matplotlib, seaborn.',
            'Завантажено датасети California Housing (регресія) та Iris (класифікація).',
            'Розділено дані на навчальну та тестову вибірки, виконано нормалізацію ознак.',
            'Побудовано регресійні моделі: LinearRegression, Ridge, Lasso, '
            'DecisionTreeRegressor, RandomForestRegressor.',
            'Побудовано класифікаційні моделі: LogisticRegression, KNeighborsClassifier, '
            'DecisionTreeClassifier, RandomForestClassifier.',
            'Обчислено метрики якості: для регресії — MAE, MSE, R²; '
            'для класифікації — accuracy, confusion matrix, classification report.',
            'Побудовано графіки якості (залишки, порівняння передбачених та справжніх значень, '
            'матриці неточностей).',
            'Здійснено порівняльний аналіз моделей за метриками.',
        ],
        'conclusion': (
            'Побудовано та оцінено регресійні та класифікаційні моделі. '
            'Для регресії (California Housing) Random Forest показав найкращий R². '
            'Для класифікації (Iris) всі моделі показали високу точність (>90%). '
            'Ансамблеві методи (Random Forest) стабільно дають найкращі результати.'
        ),
        'plots_dir': 'lab3_plots',
        'script_name': 'lab3_sklearn_models.py',
    },
    {
        'number': 4,
        'theme': 'Підбір гіперпараметрів для моделей машинного навчання',
        'goal': (
            'ознайомлення з методами підбору гіперпараметрів для моделей '
            'машинного навчання з використанням бібліотек HyperOpt та GridSearchCV, '
            'засвоєння практичних підходів до оптимізації якості моделей'
        ),
        'steps': [
            'Імпортовано бібліотеки scikit-learn, hyperopt, pandas, numpy, matplotlib.',
            'Завантажено датасети California Housing (регресія) та Iris (класифікація).',
            'Для регресії побудовано базові моделі: GradientBoostingRegressor, '
            'RandomForestRegressor, MLPRegressor.',
            'Для класифікації побудовано: RandomForestClassifier, XGBClassifier '
            '(або GradientBoostingClassifier).',
            'Для кожної моделі здійснено підбір гіперпараметрів за допомогою GridSearchCV.',
            'Для кожної моделі реалізовано оптимізацію за допомогою HyperOpt (TPE).',
            'Порівняно результати GridSearchCV та HyperOpt за метриками R², accuracy, F1.',
            'Побудовано графіки зміни метрик.',
        ],
        'conclusion': (
            'Проведено підбір гіперпараметрів двома методами. '
            'GridSearchCV забезпечує систематичний перебір, HyperOpt — стохастичну оптимізацію. '
            'Обидва методи показали порівнянні результати, але HyperOpt ефективніший '
            'при великому просторі пошуку.'
        ),
        'plots_dir': 'lab4_plots',
        'script_name': 'lab4_hyperparameters.py',
    },
    {
        'number': 5,
        'theme': 'Прогнозування/передбачення даних методами бустингу',
        'goal': (
            'навчання прогнозувати/передбачувати значення змінної '
            'за допомогою методів бустингу (Gradient Boosting, XGBoost, LightGBM)'
        ),
        'steps': [
            'Завантажено та описано датасет California Housing.',
            'Виконано підготовку даних: обробка пропущених значень, масштабування.',
            'Підготовлено тренувальний та тестовий набори.',
            'Навчено моделі бустингу: Gradient Boosting (sklearn), XGBoost, LightGBM '
            'з використанням cross-validation.',
            'Оцінено ефективність кожної моделі за метриками RMSE, MAE, R².',
            'Порівняно ефективність різних методів бустингу.',
            'Побудовано графіки порівняння метрик та feature importance.',
        ],
        'conclusion': (
            'Проведено прогнозування вартості житла методами бустингу. '
            'Усі методи показали високу якість (R² > 0.8). '
            'Найбільш інформативна ознака — MedInc (медіанний дохід). '
            'Cross-validation підтверджує стабільність моделей.'
        ),
        'plots_dir': 'lab5_plots',
        'script_name': 'lab5_boosting.py',
    },
    {
        'number': 6,
        'theme': 'Прогнозування/передбачення даних з допомогою нейронних мереж',
        'goal': (
            'навчання прогнозувати/передбачувати значення змінної '
            'за допомогою нейронних мереж (MLPRegressor різних архітектур)'
        ),
        'steps': [
            'Завантажено та описано датасет California Housing.',
            'Виконано підготовку даних, масштабування (StandardScaler — '
            'обов\'язково для нейронних мереж).',
            'Побудовано та навчено 5 архітектур MLP з різними параметрами: '
            'різна кількість шарів (1-3), різні розміри шарів, різні функції активації.',
            'Проведено cross-validation для оцінки стабільності.',
            'Оцінено ефективність кожної моделі за RMSE, MAE, R².',
            'Побудовано графіки: криві навчання (loss curves), '
            'передбачені vs справжні значення, аналіз залишків.',
        ],
        'conclusion': (
            'Протестовано 5 архітектур MLP для прогнозування. '
            'Нейронні мережі з 2-3 прихованими шарами дають кращі результати. '
            'Функція активації ReLU працює краще за tanh. '
            'Масштабування даних критичне для нейронних мереж.'
        ),
        'plots_dir': 'lab6_plots',
        'script_name': 'lab6_neural_networks.py',
    },
]


# ─────────────────────────────────────────────
#  Утиліти форматування DOCX
# ─────────────────────────────────────────────

def set_paragraph_format(paragraph, font_name="Times New Roman", font_size=14,
                         bold=False, italic=False, alignment=None,
                         line_spacing=1.5, first_indent=None,
                         space_before=None, space_after=None):
    """Налаштування форматування абзацу."""
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    pf.line_spacing = line_spacing
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)

    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        r_elem = run._element
        rPr = r_elem.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} w:cs="{font_name}" w:eastAsia="{font_name}"/>'
            )
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:cs'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph(doc, text, font_size=14, bold=False, italic=False,
                  alignment=None, line_spacing=1.5, first_indent=1.25,
                  space_before=0, space_after=0):
    """Додати абзац з форматуванням."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_paragraph_format(
        p, font_size=font_size, bold=bold, italic=italic,
        alignment=alignment, line_spacing=line_spacing,
        first_indent=first_indent, space_before=space_before,
        space_after=space_after
    )
    return p


def add_empty_paragraph(doc, count=1):
    """Додати порожній абзац."""
    for _ in range(count):
        p = doc.add_paragraph()
        run = p.add_run("")
        set_paragraph_format(p, line_spacing=1.0, first_indent=None)


def add_page_break(doc):
    """Додати розрив сторінки."""
    p = doc.add_paragraph()
    run = p.add_run()
    run_elem = run._element
    br_elem = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run_elem.append(br_elem)


def add_image_centered(doc, img_path, width_cm=14.0):
    """Додати зображення по центру."""
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        return True
    return False


def add_image_caption(doc, number, caption):
    """Підпис рисунка під зображенням."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"Рисунок {number} – {caption}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    r_elem = run._element
    rPr = r_elem.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:cs="Times New Roman" w:eastAsia="Times New Roman"/>'
    )
    rPr.insert(0, rFonts)


# ─────────────────────────────────────────────
#  Генерація одного звіту
# ─────────────────────────────────────────────

def generate_lab_report(lab):
    """Згенерувати DOCX-звіт для однієї лабораторної роботи."""
    doc = Document()

    # Налаштування сторінки
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.0)

    # Стиль Normal
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5

    lab_num = lab['number']
    lab_theme = lab['theme']

    # ══════════════════════════════════════════
    #  ТИТУЛЬНА СТОРІНКА
    # ══════════════════════════════════════════

    add_paragraph(doc,
        "Міністерство освіти і науки України",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        "Вінницький національний технічний університет",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        "Факультет інтелектуальних інформаційних технологій та автоматизації",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )

    add_empty_paragraph(doc, 1)

    add_paragraph(doc,
        f"Кафедра {DEPARTMENT}",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )

    add_empty_paragraph(doc, 3)

    add_paragraph(doc,
        f"Лабораторна робота № {lab_num}",
        font_size=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.5
    )

    add_empty_paragraph(doc, 1)

    add_paragraph(doc,
        f"з дисципліни «{DISCIPLINE}»",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.5
    )
    add_paragraph(doc,
        f"з теми: «{lab_theme}»",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.5
    )

    add_empty_paragraph(doc, 5)

    # Виконав / Перевірив
    add_paragraph(doc,
        f"Виконав: студент групи {STUDENT_GROUP}",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, line_spacing=1.0, space_after=0
    )
    add_paragraph(doc,
        f"{STUDENT_FULL_NAME_CAPS}",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, line_spacing=1.0, space_after=6
    )
    add_paragraph(doc,
        f"Перевірив: {SUPERVISOR_TITLE}",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, line_spacing=1.0, space_after=0
    )
    add_paragraph(doc,
        f"{SUPERVISOR_DEGREE} {SUPERVISOR_NAME}",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, line_spacing=1.0, space_after=0
    )

    add_empty_paragraph(doc, 5)

    add_paragraph(doc,
        f"{CITY} {UNIVERSITY_SHORT} – {YEAR} року",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, line_spacing=1.0
    )

    # ══════════════════════════════════════════
    #  ОСНОВНА ЧАСТИНА
    # ══════════════════════════════════════════
    add_page_break(doc)

    # Мета роботи
    add_paragraph(doc,
        f"Метою роботи є {lab['goal']}.",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_indent=1.25, space_after=6
    )

    # Хід роботи
    add_paragraph(doc,
        "Хід роботи:",
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=1.25, space_after=6
    )

    for i, step in enumerate(lab['steps'], 1):
        add_paragraph(doc,
            f"{i}. {step}",
            font_size=14,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_indent=1.25, space_after=3
        )

    # Вставка графіків (якщо існують)
    plots_dir = os.path.join(SCRIPT_DIR, lab['plots_dir'])
    if os.path.exists(plots_dir):
        add_empty_paragraph(doc, 1)
        add_paragraph(doc,
            "Результати виконання роботи:",
            font_size=14, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_indent=1.25, space_after=6
        )

        png_files = sorted([f for f in os.listdir(plots_dir) if f.endswith('.png')])
        for fig_num, png_file in enumerate(png_files, 1):
            img_path = os.path.join(plots_dir, png_file)
            # Назва з файлу (без .png, замінюємо _ на пробіли)
            caption = png_file.replace('.png', '').replace('_', ' ').capitalize()
            added = add_image_centered(doc, img_path, width_cm=15.0)
            if added:
                add_image_caption(doc, f"{lab_num}.{fig_num}", caption)

    # Посилання на ноутбук
    add_empty_paragraph(doc, 1)
    add_paragraph(doc,
        f"Результат виконання роботи наведено у скрипті: {lab['script_name']}",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_indent=1.25, space_after=6
    )

    # Висновок
    add_paragraph(doc,
        f"Висновок. {lab['conclusion']}",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_indent=1.25, space_after=6
    )

    # Збереження
    filename = f"Лабораторна_робота_{lab_num}.docx"
    output_path = os.path.join(SCRIPT_DIR, filename)
    doc.save(output_path)
    return output_path, filename


# ─────────────────────────────────────────────
#  Головна функція
# ─────────────────────────────────────────────

def main():
    print("=" * 60)
    print("Генератор звітів до лабораторних робіт")
    print(f"Дисципліна: {DISCIPLINE}")
    print(f"Студент: {STUDENT_FULL_NAME}, група {STUDENT_GROUP}")
    print("=" * 60)

    for lab in LABS:
        print(f"\nГенерація звіту: Лабораторна робота №{lab['number']}...")
        output_path, filename = generate_lab_report(lab)
        print(f"  ✓ Збережено: {filename}")

    print(f"\n{'=' * 60}")
    print(f"Готово! Згенеровано {len(LABS)} звітів.")
    print(f"{'=' * 60}")


if __name__ == '__main__':
    main()
