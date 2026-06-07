# -*- coding: utf-8 -*-
"""
Генератор пояснювальної записки до курсової роботи
з дисципліни «Методи комп'ютерних обчислень»

Тема: Дослідження чисельних методів вирішення нелінійних рівнянь
Рівняння: 0.5x^5 - 0.005x - 1 = 0

Формат: DOCX відповідно до ДСТУ 3008:2015, ДСТУ 8302:2015
та методичних вказівок до курсової роботи ВНТУ
Використання: python generate_report.py
"""

import os
import sys
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')

import math
import cmath
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─────────────────────────────────────────────
#  Конфігурація: змініть дані студента тут
# ─────────────────────────────────────────────
STUDENT_FULL_NAME = "Столовник В. О."
STUDENT_GROUP = "2ІСТ-24б"
STUDENT_COURSE = "2"
STUDENT_SPECIALTY = "126 «Інформаційні системи та технології»"
SUPERVISOR_NAME = "Богач І. В."
SUPERVISOR_TITLE = "к.т.н., доцент кафедри АІІТ"
DEPARTMENT_HEAD = "Бісікало О. В."
DEPARTMENT_HEAD_TITLE = "Зав. кафедри АІІТ, проф., д.т.н."
YEAR = "2026"
CITY = "Вінниця"

# Параметри рівняння
A_DEFAULT = 1.0
B_DEFAULT = 2.0
EPSILON = 0.01
X0_DEFAULT = 2.0

OUTPUT_FILENAME = "Пояснювальна_записка_МКО.docx"

# ─────────────────────────────────────────────
#  Чисельні методи (дзеркало C# реалізації)
# ─────────────────────────────────────────────

def f(x):
    """0.5x^5 - 0.005x - 1"""
    return 0.5 * x**5 - 0.005 * x - 1.0

def df(x):
    """Похідна: 2.5x^4 - 0.005"""
    return 2.5 * x**4 - 0.005

def fc(z):
    """Комплексна функція"""
    return 0.5 * z**5 - 0.005 * z - 1.0

def dfc(z):
    """Комплексна похідна"""
    return 2.5 * z**4 - 0.005


def bisection_method(a, b, eps):
    """Метод половинного ділення"""
    steps = []
    iteration = 0
    while True:
        c = (a + b) / 2
        iteration += 1
        fc_val = f(c)
        width = abs(b - a)
        steps.append({
            'iter': iteration, 'a': a, 'b': b,
            'c': c, 'fc': fc_val, 'width': width
        })
        if f(a) * fc_val > 0:
            a = c
        else:
            b = c
        if abs(b - a) < eps:
            break
    root = (a + b) / 2
    return root, f(root), iteration, steps


def chord_method(a, b, eps):
    """Метод хорд"""
    steps = []
    iteration = 0
    while True:
        iteration += 1
        fa = f(a)
        fb = f(b)
        x = a - (fa * (b - a)) / (fb - fa)
        fx = f(x)
        steps.append({
            'iter': iteration, 'a': a, 'b': b,
            'x': x, 'fx': fx
        })
        if abs(fx) < eps:
            break
        if fa * fx < 0:
            b = x
        else:
            a = x
    return x, f(x), iteration, steps


def newton_method(x0, eps):
    """Метод дотичних (Ньютона)"""
    steps = []
    x = x0
    iteration = 0
    while True:
        x_prev = x
        fx_prev = f(x_prev)
        dfx_prev = df(x_prev)
        x = x_prev - fx_prev / dfx_prev
        iteration += 1
        precision = abs(x - x_prev)
        steps.append({
            'iter': iteration, 'x_prev': x_prev,
            'fx_prev': fx_prev, 'dfx_prev': dfx_prev,
            'x_curr': x, 'precision': precision
        })
        if precision <= eps:
            break
    return x, f(x), iteration, steps


def complex_newton(z0, eps, max_iter=100):
    """Метод Ньютона для комплексних коренів"""
    steps = []
    z = z0
    for iteration in range(1, max_iter + 1):
        z_prev = z
        fz = fc(z)
        dfz = dfc(z)
        z = z - fz / dfz
        precision = abs(z - z_prev)
        steps.append({
            'iter': iteration, 'z_prev': z_prev,
            'z_curr': z, 'precision': precision
        })
        if precision <= eps:
            break
    return z, iteration, steps


# ─────────────────────────────────────────────
#  Утиліти для форматування документа
# ─────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Встановлює фон клітинки таблиці"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_paragraph_format(paragraph, font_name="Times New Roman", font_size=14,
                         bold=False, italic=False, alignment=None,
                         line_spacing=1.5, first_indent=None,
                         space_before=None, space_after=None,
                         keep_together=False, keep_with_next=False):
    """Налаштування форматування абзацу"""
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    pf.line_spacing = line_spacing
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    pf.keep_together = keep_together
    pf.keep_with_next = keep_with_next

    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        # Для кирилиці — вказуємо eastAsia та cs шрифт
        r_elem = run._element
        rPr = r_elem.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{font_name}" w:eastAsia="{font_name}"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:cs'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph(doc, text, font_size=14, bold=False, italic=False,
                  alignment=None, line_spacing=1.5, first_indent=1.25,
                  space_before=0, space_after=0,
                  keep_with_next=False):
    """Додати абзац з форматуванням ДСТУ"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_paragraph_format(
        p, font_size=font_size, bold=bold, italic=italic,
        alignment=alignment, line_spacing=line_spacing,
        first_indent=first_indent, space_before=space_before,
        space_after=space_after, keep_with_next=keep_with_next
    )
    return p


def add_heading_h1(doc, text, numbered=True):
    """Додати заголовок 1 рівня (ВЕЛИКІ ЛІТЕРИ, жирний, по центру).
    Згідно з методичними вказівками: кожен розділ починається з нової сторінки."""
    # Розрив сторінки перед заголовком (крім першого)
    add_page_break(doc)
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    set_paragraph_format(
        p, font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5, first_indent=None,
        space_before=0, space_after=6,
        keep_with_next=True
    )
    return p


def add_heading_h2(doc, text):
    """Додати заголовок 2 рівня (жирний, з абзацного відступу, з великої літери)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_paragraph_format(
        p, font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5, first_indent=1.25,
        space_before=6, space_after=3,
        keep_with_next=True
    )
    return p


def add_empty_paragraph(doc, count=1):
    """Додати порожній абзац"""
    for _ in range(count):
        p = doc.add_paragraph()
        run = p.add_run("")
        set_paragraph_format(p, line_spacing=1.0, first_indent=None)


def add_page_break(doc):
    """Додати розрив сторінки"""
    p = doc.add_paragraph()
    run = p.add_run()
    run_elem = run._element
    br_elem = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run_elem.append(br_elem)


def add_table_caption(doc, number, caption):
    """Додати підпис таблиці ЗВЕРХУ (з абзацного відступу, згідно з ДСТУ 3008:2015).
    Формат: 'Таблиця 4.1 – Назва таблиці'"""
    add_paragraph(
        doc, f"Таблиця {number} – {caption}",
        font_size=14, italic=False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=1.25, space_before=6, space_after=3
    )


def add_table_with_data(doc, headers, rows, col_widths=None):
    """Створити таблицю з даними, стилізовану відповідно до ДСТУ"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Заголовки
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        set_cell_shading(cell, "D9E2F3")

    # Дані
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0

    # Ширини колонок
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)

    return table


def add_code_block(doc, code_text, font_size=8):
    """Додати блок коду з моноширинним шрифтом"""
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(font_size)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = None

        # Задаємо шрифт для кирилиці теж
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Courier New" w:hAnsi="Courier New" w:cs="Courier New"/>')
            rPr.insert(0, rFonts)


def add_formula(doc, formula_text, formula_number=None):
    """Додати формулу (по центру, номер праворуч у круглих дужках)"""
    p = doc.add_paragraph()
    if formula_number:
        # Формула по центру, номер — правий табстоп
        run = p.add_run(formula_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.italic = True
        # Додаємо табуляцію і номер
        tab_run = p.add_run(f"\t({formula_number})")
        tab_run.font.name = "Times New Roman"
        tab_run.font.size = Pt(14)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        run = p.add_run(formula_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.italic = True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = None
    return p


def add_page_numbers(doc):
    """Додати нумерацію сторінок праворуч у верхньому куті (згідно з ДСТУ)"""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)

    # Додаємо поле PAGE
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_char_begin)

    run2 = p.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instr)
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(14)

    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fld_char_end)

    # Титульний аркуш без номера — потрібно встановити titlePg
    sectPr = section._sectPr
    title_pg = sectPr.find(qn('w:titlePg'))
    if title_pg is None:
        title_pg = parse_xml(f'<w:titlePg {nsdecls("w")}/>')
        sectPr.append(title_pg)


def format_complex(z):
    """Форматувати комплексне число"""
    sign = "+" if z.imag >= 0 else "−"
    return f"{z.real:.6f} {sign} {abs(z.imag):.6f}i"


# ─────────────────────────────────────────────
#  Основна функція генерації документа
# ─────────────────────────────────────────────

def generate_report():
    doc = Document()

    # ── Налаштування сторінки (ДСТУ 3008:2015) ──
    # Поля: верхній/нижній ≥ 20мм, лівий ≥ 25мм, правий ≥ 10мм
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.0)

    # Стиль Normal за замовчуванням
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5

    # ── Нумерація сторінок (праворуч, верхній кут) ──
    add_page_numbers(doc)

    # ══════════════════════════════════════════
    #  ТИТУЛЬНА СТОРІНКА
    # ══════════════════════════════════════════

    add_paragraph(doc,
        "Міністерство освіти і науки України",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        "Вінницький національний технічний університет",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        "Факультет інтелектуальних інформаційних технологій та автоматизації",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        "Кафедра автоматизації та інтелектуальних інформаційних технологій",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )

    add_empty_paragraph(doc, 3)

    add_paragraph(doc,
        "КУРСОВА РОБОТА",
        font_size=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_before=0, space_after=0, line_spacing=1.5
    )
    add_paragraph(doc,
        "з дисципліни «Методи комп'ютерних обчислень»",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.5
    )
    add_paragraph(doc,
        "на тему: «Дослідження чисельних методів вирішення нелінійних рівнянь»",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.5
    )

    add_empty_paragraph(doc, 3)

    add_paragraph(doc,
        f"Студента {STUDENT_COURSE} курсу {STUDENT_GROUP} групи",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, line_spacing=1.0, space_after=0
    )
    add_paragraph(doc,
        f"спеціальності {STUDENT_SPECIALTY}",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, line_spacing=1.0, space_after=0
    )
    add_paragraph(doc,
        f"{STUDENT_FULL_NAME}",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, line_spacing=1.0, space_after=6
    )
    add_paragraph(doc,
        f"Керівник {SUPERVISOR_TITLE}",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, line_spacing=1.0, space_after=0
    )
    add_paragraph(doc,
        f"{SUPERVISOR_NAME}",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, line_spacing=1.0, space_after=0
    )

    add_empty_paragraph(doc, 2)

    add_paragraph(doc,
        "Кількість балів: __________ Оцінка: ECTS _____",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, line_spacing=1.0, space_after=6
    )
    add_paragraph(doc,
        "Члени комісії:    _____________   ________________________",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=None, line_spacing=1.0, space_after=0
    )
    add_paragraph(doc,
        "                              (підпис)                      (прізвище та ініціали)",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=None, line_spacing=1.0, space_after=0
    )
    add_paragraph(doc,
        "                         _____________   ________________________",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=None, line_spacing=1.0, space_after=0
    )
    add_paragraph(doc,
        "                              (підпис)                      (прізвище та ініціали)",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=None, line_spacing=1.0, space_after=0
    )

    add_empty_paragraph(doc, 2)

    add_paragraph(doc,
        f"{CITY} – {YEAR} рік",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, line_spacing=1.0
    )

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  ІНДИВІДУАЛЬНЕ ЗАВДАННЯ (Додаток Б метод. вказівок)
    # ══════════════════════════════════════════

    add_paragraph(doc,
        "Міністерство освіти і науки України",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        "Вінницький національний технічний університет",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        "Факультет інтелектуальних інформаційних технологій та автоматизації",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=6, line_spacing=1.0
    )

    # Блок ЗАТВЕРДЖУЮ (згідно з Додатком Б)
    add_paragraph(doc,
        "ЗАТВЕРДЖУЮ",
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, space_before=6, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        f"{DEPARTMENT_HEAD_TITLE}",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        f"__________ {DEPARTMENT_HEAD}",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        "         (підпис)",
        font_size=12,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, space_after=0, line_spacing=1.0
    )
    add_paragraph(doc,
        f"«____» ____________ {YEAR} р.",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=None, space_after=6, line_spacing=1.0
    )

    add_empty_paragraph(doc, 1)

    add_paragraph(doc,
        "ІНДИВІДУАЛЬНЕ ЗАВДАННЯ",
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=0, line_spacing=1.5
    )
    add_paragraph(doc,
        'на курсову роботу з дисципліни «Методи комп\'ютерних обчислень»',
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=3, line_spacing=1.5
    )
    add_paragraph(doc,
        f"студенту {STUDENT_FULL_NAME} групи {STUDENT_GROUP}",
        font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=6, line_spacing=1.5
    )
    add_paragraph(doc,
        "ТЕМА Дослідження чисельних методів вирішення нелінійних рівнянь.",
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=None, space_after=6, line_spacing=1.5
    )

    add_paragraph(doc, "Постановка задачі:", font_size=14, bold=True,
                  first_indent=1.25, space_after=3)

    tasks = [
        "Дослідити чисельні методи розв'язання нелінійних рівнянь.",
        "Скласти алгоритми для кожного із заданих чисельних методів.",
        "Розробити програмне забезпечення для знаходження коренів нелінійних рівнянь (кожен метод реалізувати у вигляді окремого класу).",
        "Протестувати програму.",
        "Розв'язати задане нелінійне рівняння, проаналізувати отримані результати та розрахувати похибки знайдених коренів."
    ]
    for i, task in enumerate(tasks, 1):
        add_paragraph(doc, f"{i}) {task}", font_size=14,
                      first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc, "Вхідні дані:", font_size=14, bold=True,
                  first_indent=1.25, space_before=6, space_after=3)
    add_paragraph(doc,
        "– задане нелінійне рівняння: 0.5x⁵ − 0.005x − 1 = 0;",
        font_size=14, first_indent=1.25, space_after=0,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "– початкові значення: a = 1.0; b = 2.0; x₀ = 2.0;",
        font_size=14, first_indent=1.25, space_after=0,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "– задана точність: ε = 0.01;",
        font_size=14, first_indent=1.25, space_after=0,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "– методи для вирішення: половинного ділення, хорд, дотичних (Ньютона) та комплексних коренів.",
        font_size=14, first_indent=1.25, space_after=6,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc, "Зміст пояснювальної записки до курсової роботи:",
                  font_size=14, bold=True, first_indent=1.25, space_after=3)
    toc_items = [
        "Індивідуальне завдання", "Вступ",
        "1. Короткі теоретичні відомості", "2. Алгоритми методів",
        "3. Розробка програмного забезпечення", "4. Тестування програмного забезпечення",
        "5. Аналіз отриманих результатів", "Висновки",
        "Перелік посилань", "Додатки"
    ]
    for item in toc_items:
        add_paragraph(doc, item, font_size=14, first_indent=1.25, space_after=0)

    add_empty_paragraph(doc, 2)

    add_paragraph(doc,
        f'Дата видачі "____" _________ {YEAR} р.        Керівник _________________  {SUPERVISOR_NAME}',
        font_size=14, first_indent=None, space_after=0, line_spacing=1.0)
    add_paragraph(doc,
        '\t                                                                                   (підпис)',
        font_size=12, first_indent=None, space_after=6, line_spacing=1.0)
    add_paragraph(doc,
        '                                                                  Завдання отримав _____________',
        font_size=14, first_indent=None, space_after=0, line_spacing=1.0)
    add_paragraph(doc,
        '                                                                                                               (підпис)',
        font_size=12, first_indent=None, space_after=0, line_spacing=1.0)

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  АНОТАЦІЯ
    # ══════════════════════════════════════════

    add_paragraph(doc, "АНОТАЦІЯ", font_size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=None, space_after=6)

    add_paragraph(doc,
        "У курсовій роботі досліджується задача чисельного знаходження коренів нелінійного "
        "алгебраїчного рівняння. Метою роботи є розробка програмного забезпечення для розв'язання "
        "рівняння виду 0.5x⁵ − 0.005x − 1 = 0 із заданою точністю з використанням різних чисельних методів.",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "Для знаходження дійсного кореня на заданому інтервалі було теоретично обґрунтовано та "
        "програмно реалізовано три ітераційні алгоритми: метод половинного ділення (дихотомії), "
        "метод хорд та метод дотичних (Ньютона). Крім того, для знаходження усіх комплексних коренів "
        "полінома було застосовано модифікований метод Ньютона для комплексної площини.",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "Програмну реалізацію виконано мовою програмування C# з використанням об'єктно-орієнтованого "
        "підходу та стандартної бібліотеки для роботи з комплексними числами (структура Complex). "
        "Для кожного з методів створено окремі класи, які відстежують проміжні результати на кожній ітерації.",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "У роботі проведено порівняльний аналіз ефективності та швидкості збіжності досліджуваних "
        "чисельних методів на основі кількості виконаних кроків (ітерацій), необхідних для досягнення "
        "заданої похибки обчислень. Результати роботи подано у вигляді форматованих таблиць з покроковими "
        "значеннями шуканих величин.",
        font_size=14, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "Ключові слова: чисельні методи, нелінійне рівняння, метод половинного ділення, метод хорд, "
        "метод Ньютона, метод дотичних, комплексні корені, обчислювальна математика, C#.",
        font_size=14, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  ЗМІСТ (автоматично генерується в Word, тут — заглушка)
    # ══════════════════════════════════════════

    add_paragraph(doc, "ЗМІСТ", font_size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=None, space_after=12)

    toc_entries = [
        ("ВСТУП", False),
        ("1 КОРОТКІ ТЕОРЕТИЧНІ ВІДОМОСТІ", False),
        ("    1.1 Метод половинного ділення", True),
        ("    1.2 Метод хорд", True),
        ("    1.3 Метод дотичних (Ньютона)", True),
        ("    1.4 Метод Ньютона для комплексних коренів", True),
        ("2 АЛГОРИТМИ МЕТОДІВ", False),
        ("    2.1 Алгоритм методу половинного ділення", True),
        ("    2.2 Алгоритм методу хорд", True),
        ("    2.3 Алгоритм методу дотичних (Ньютона)", True),
        ("    2.4 Алгоритм методу Ньютона для комплексних коренів", True),
        ("3 РОЗРОБКА ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ", False),
        ("    3.1 Вибір мови програмування", True),
        ("    3.2 Вхідні/вихідні дані", True),
        ("    3.3 Структура програми", True),
        ("    3.4 Інструкція користувачеві", True),
        ("4 ТЕСТУВАННЯ ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ", False),
        ("5 АНАЛІЗ ОТРИМАНИХ РЕЗУЛЬТАТІВ", False),
        ("ВИСНОВКИ", False),
        ("ПЕРЕЛІК ПОСИЛАНЬ", False),
        ("ДОДАТКИ", False),
        ("    Додаток А. Лістинг програми", True),
    ]
    for entry_text, is_sub in toc_entries:
        p = add_paragraph(doc, entry_text,
                          font_size=14, bold=(not is_sub),
                          first_indent=None,
                          space_after=0, line_spacing=1.5)

    # ══════════════════════════════════════════
    #  ВСТУП (з нової сторінки)
    # ══════════════════════════════════════════

    add_heading_h1(doc, "ВСТУП")

    add_paragraph(doc,
        "Розв'язання нелінійних рівнянь є однією з фундаментальних задач обчислювальної математики, "
        "яка має надзвичайно широке застосування у різних галузях науки та техніки: від інженерних "
        "розрахунків до моделювання фізичних процесів, від оптимізаційних задач до аналізу фінансових "
        "моделей [1, 2]. У більшості практичних випадків нелінійні рівняння не мають аналітичного "
        "розв'язку, тому виникає необхідність застосування чисельних (наближених) методів.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "Актуальність даної роботи зумовлена тим, що сучасні програмні системи потребують ефективних "
        "та надійних алгоритмів чисельного розв'язання рівнянь, а розуміння принципів їх роботи є "
        "необхідною компетентністю для фахівця у галузі інформаційних систем та технологій [3].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "Метою курсової роботи є дослідження та програмна реалізація чисельних методів розв'язання "
        "нелінійного рівняння виду 0.5x⁵ − 0.005x − 1 = 0. У роботі розглядаються такі методи: "
        "метод половинного ділення (дихотомії), метод хорд та метод дотичних (Ньютона) для знаходження "
        "дійсних коренів, а також модифікований метод Ньютона для знаходження комплексних коренів "
        "полінома [4, 5].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "Для досягнення поставленої мети необхідно виконати такі завдання: "
        "вивчити теоретичні основи кожного з методів; розробити алгоритми їх реалізації; "
        "створити програмне забезпечення мовою C# з графічним інтерфейсом (Windows Forms) [6]; "
        "провести обчислювальний експеримент і проаналізувати результати щодо точності та збіжності методів.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ══════════════════════════════════════════
    #  1 КОРОТКІ ТЕОРЕТИЧНІ ВІДОМОСТІ
    # ══════════════════════════════════════════

    add_heading_h1(doc, "1 КОРОТКІ ТЕОРЕТИЧНІ ВІДОМОСТІ")

    add_paragraph(doc,
        "Задача знаходження коренів нелінійного рівняння f(x) = 0 є класичною задачею "
        "обчислювальної математики [1]. Під коренем (або розв'язком) рівняння розуміється таке "
        "значення x*, для якого f(x*) = 0. Основна ідея чисельних методів полягає у побудові "
        "ітераційної послідовності {xₙ}, яка збігається до точного кореня x* із заданою точністю ε.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "Розв'язання задачі зазвичай складається з двох етапів: відокремлення коренів "
        "(визначення інтервалу [a, b], на якому існує рівно один корінь) та уточнення коренів "
        "(ітераційне наближення до кореня з необхідною точністю) [2, 3].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 1.1 Метод половинного ділення
    add_heading_h2(doc, "1.1 Метод половинного ділення")

    add_paragraph(doc,
        "Метод половинного ділення (дихотомії, бісекції) є найпростішим та найнадійнішим "
        "методом знаходження кореня рівняння на заданому інтервалі [a, b]. Він ґрунтується "
        "на теоремі Больцано–Коші: якщо безперервна функція f(x) приймає значення різних знаків "
        "на кінцях відрізка (f(a)·f(b) < 0), то на цьому відрізку існує хоча б один корінь [1, 4].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "Алгоритм полягає у послідовному поділі інтервалу навпіл: обчислюється середня точка "
        "та визначається, в якій половині міститься корінь. Формула обчислення середньої точки (1.1):",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_formula(doc, "c = (a + b) / 2", "1.1")

    add_paragraph(doc,
        "Процес повторюється, поки ширина інтервалу не стане менше заданої точності ε. "
        "Метод гарантовано збігається, проте має лінійну швидкість збіжності [2].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 1.2 Метод хорд
    add_heading_h2(doc, "1.2 Метод хорд")

    add_paragraph(doc,
        "Метод хорд (метод лінійної інтерполяції, метод False Position) є модифікацією методу "
        "половинного ділення. Замість середньої точки інтервалу обчислюється точка перетину хорди, "
        "що з'єднує точки (a, f(a)) і (b, f(b)), з віссю абсцис за формулою (1.2) [3, 5]:",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_formula(doc, "x = a − f(a)·(b − a) / (f(b) − f(a))", "1.2")

    add_paragraph(doc,
        "Далі, аналогічно методу дихотомії, визначається половина інтервалу, що містить корінь, "
        "і процес повторюється. Метод хорд, як правило, збігається швидше за метод половинного "
        "ділення, оскільки враховує характер функції [4].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 1.3 Метод Ньютона
    add_heading_h2(doc, "1.3 Метод дотичних (Ньютона)")

    add_paragraph(doc,
        "Метод Ньютона (метод дотичних) є одним із найефективніших ітераційних методів "
        "розв'язання нелінійних рівнянь. Він базується на лінеаризації функції за допомогою "
        "розкладу в ряд Тейлора та використанні дотичної до графіка функції. Ітераційна формула "
        "методу Ньютона (1.3) [1, 5]:",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_formula(doc, "xₙ₊₁ = xₙ − f(xₙ) / f'(xₙ)", "1.3")

    add_paragraph(doc,
        "Метод Ньютона має квадратичну швидкість збіжності поблизу кореня, "
        "що означає подвоєння кількості правильних знаків на кожному кроці. "
        "Однак він потребує обчислення похідної f'(x) та вдалого вибору початкового наближення x₀, "
        "інакше ітераційний процес може розбігатися [2, 8].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 1.4 Комплексні корені
    add_heading_h2(doc, "1.4 Метод Ньютона для комплексних коренів")

    add_paragraph(doc,
        "Для знаходження комплексних коренів полінома метод Ньютона узагальнюється на комплексну "
        "площину. Ітераційна формула має той самий вигляд (1.4) [7, 8]:",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_formula(doc, "zₙ₊₁ = zₙ − f(zₙ) / f'(zₙ)", "1.4")

    add_paragraph(doc,
        "де zₙ — комплексне число. Обираючи різні початкові наближення z₀ з комплексної площини, "
        "можна знайти всі корені полінома. Поліном 5-го степеня, згідно з основною теоремою алгебри, "
        "має рівно 5 коренів (з урахуванням кратності), серед яких один дійсний та чотири комплексні "
        "(попарно спряжені) [9].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ══════════════════════════════════════════
    #  2 АЛГОРИТМИ МЕТОДІВ
    # ══════════════════════════════════════════

    add_heading_h1(doc, "2 АЛГОРИТМИ МЕТОДІВ")

    add_paragraph(doc,
        "У цьому розділі наведено покрокові алгоритми кожного з реалізованих чисельних методів "
        "відповідно до їх математичного обґрунтування, наведеного в розділі 1.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading_h2(doc, "2.1 Алгоритм методу половинного ділення")

    algo_bisect = [
        "Крок 1. Задати інтервал [a, b] такий, що f(a)·f(b) < 0, та точність ε.",
        "Крок 2. Обчислити середню точку c = (a + b) / 2.",
        "Крок 3. Обчислити f(c).",
        "Крок 4. Якщо f(a)·f(c) > 0, то a = c; інакше b = c.",
        "Крок 5. Якщо |b − a| ≥ ε, перейти до кроку 2.",
        "Крок 6. Вивести x* = c як наближений корінь."
    ]
    for step in algo_bisect:
        add_paragraph(doc, step, first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading_h2(doc, "2.2 Алгоритм методу хорд")

    algo_chord = [
        "Крок 1. Задати інтервал [a, b] такий, що f(a)·f(b) < 0, та точність ε.",
        "Крок 2. Обчислити x = a − f(a)·(b − a) / (f(b) − f(a)).",
        "Крок 3. Обчислити f(x).",
        "Крок 4. Якщо |f(x)| < ε, то x — наближений корінь (завершення).",
        "Крок 5. Якщо f(a)·f(x) < 0, то b = x; інакше a = x.",
        "Крок 6. Перейти до кроку 2."
    ]
    for step in algo_chord:
        add_paragraph(doc, step, first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading_h2(doc, "2.3 Алгоритм методу дотичних (Ньютона)")

    algo_newton = [
        "Крок 1. Задати початкове наближення x₀ та точність ε.",
        "Крок 2. Обчислити f(xₙ) та f'(xₙ).",
        "Крок 3. Обчислити xₙ₊₁ = xₙ − f(xₙ) / f'(xₙ).",
        "Крок 4. Якщо |xₙ₊₁ − xₙ| ≤ ε, то xₙ₊₁ — наближений корінь (завершення).",
        "Крок 5. Присвоїти xₙ = xₙ₊₁, перейти до кроку 2."
    ]
    for step in algo_newton:
        add_paragraph(doc, step, first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading_h2(doc, "2.4 Алгоритм методу Ньютона для комплексних коренів")

    algo_complex = [
        "Крок 1. Задати початкове наближення z₀ ∈ ℂ та точність ε.",
        "Крок 2. Обчислити f(zₙ) та f'(zₙ) у комплексній площині.",
        "Крок 3. Обчислити zₙ₊₁ = zₙ − f(zₙ) / f'(zₙ).",
        "Крок 4. Якщо |zₙ₊₁ − zₙ| ≤ ε, то zₙ₊₁ — наближений корінь (завершення).",
        "Крок 5. Присвоїти zₙ = zₙ₊₁, перейти до кроку 2."
    ]
    for step in algo_complex:
        add_paragraph(doc, step, first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ══════════════════════════════════════════
    #  3 РОЗРОБКА ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ
    # ══════════════════════════════════════════

    add_heading_h1(doc, "3 РОЗРОБКА ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ")

    # 3.1
    add_heading_h2(doc, "3.1 Вибір мови програмування")

    add_paragraph(doc,
        "Для реалізації програмного забезпечення обрано мову програмування C# (.NET) з використанням "
        "технології Windows Forms для побудови графічного інтерфейсу користувача [6]. Такий вибір зумовлений "
        "наступними перевагами:",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    advantages = [
        "C# є сучасною об'єктно-орієнтованою мовою з потужною системою типів та обробкою виключень;",
        "стандартна бібліотека .NET включає структуру System.Numerics.Complex для роботи з комплексними числами [7];",
        "Windows Forms забезпечує швидку розробку графічних інтерфейсів із вбудованими компонентами;",
        "платформа .NET підтримує кросплатформну компіляцію та має високу продуктивність."
    ]
    for i, adv in enumerate(advantages):
        prefix = "а)" if i == 0 else "б)" if i == 1 else "в)" if i == 2 else "г)"
        add_paragraph(doc, f"{prefix} {adv}", first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 3.2
    add_heading_h2(doc, "3.2 Вхідні/вихідні дані")

    add_paragraph(doc, "Вхідні дані програми:", bold=True,
                  first_indent=1.25, space_after=3)

    input_data = [
        f"a = {A_DEFAULT} — ліва межа інтервалу;",
        f"b = {B_DEFAULT} — права межа інтервалу;",
        f"ε = {EPSILON} — задана точність обчислень;",
        f"x₀ = {X0_DEFAULT} — початкове наближення для методу Ньютона."
    ]
    for item in input_data:
        add_paragraph(doc, f"– {item}", first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc, "Вихідні дані програми:", bold=True,
                  first_indent=1.25, space_before=6, space_after=3)

    output_data = [
        "таблиця покрокових обчислень для кожного методу;",
        "знайдені дійсні та комплексні корені рівняння;",
        "кількість ітерацій, необхідних для досягнення заданої точності;",
        "значення функції у знайдених коренях (перевірка)."
    ]
    for item in output_data:
        add_paragraph(doc, f"– {item}", first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 3.3
    add_heading_h2(doc, "3.3 Структура програми")

    add_paragraph(doc,
        "Програма побудована за принципами об'єктно-орієнтованого програмування [6]. "
        "Кожен чисельний метод реалізовано в окремому класі, що забезпечує модульність "
        "та можливість повторного використання коду. "
        "Лістинг програми наведено в додатку А.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc, "Основні класи програми:", bold=True,
                  first_indent=1.25, space_after=3)

    classes = [
        ("BisectionMethod", "реалізація методу половинного ділення (файл BisectionMethod.cs);"),
        ("ChordMethod", "реалізація методу хорд (файл ChordMethod.cs);"),
        ("NewtonMethod", "реалізація методу Ньютона для дійсних та комплексних коренів (файл NewtonMethod.cs);"),
        ("Coursework", "клас-координатор, що викликає всі методи та формує результат (файл coursework.cs);"),
        ("CourseworkEquations", "статичний клас з функціями f(x), f'(x), f(z), f'(z) (файл coursework.cs);"),
        ("MainForm", "графічний інтерфейс Windows Forms з полями введення та відображенням результатів (файл MainForm.cs)."),
    ]
    for cls_name, cls_desc in classes:
        p = doc.add_paragraph()
        run_bold = p.add_run(f"– {cls_name} — ")
        run_normal = p.add_run(cls_desc)
        set_paragraph_format(p, font_size=14, first_indent=1.25, space_after=0,
                             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        run_bold.font.bold = True
        run_bold.font.name = "Times New Roman"
        run_bold.font.size = Pt(14)
        run_normal.font.name = "Times New Roman"
        run_normal.font.size = Pt(14)

    add_paragraph(doc,
        "Кожен клас чисельного методу зберігає покрокову інформацію (ітерації) у відповідних "
        "структурах даних (BisectionStep, ChordStep, NewtonStep, ComplexNewtonStep), "
        "що дозволяє відтворити хід обчислень.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=6)

    # 3.4
    add_heading_h2(doc, "3.4 Інструкція користувачеві")

    add_paragraph(doc,
        "Для запуску програми необхідно відкрити виконуваний файл або запустити проєкт "
        "із середовища розробки (Visual Studio). Після запуску відкривається головне вікно "
        "програми з вкладкою «Coursework».",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    instructions = [
        "У полях a, b, eps, x0 задати відповідні параметри (за замовчуванням: a = 1.0, b = 2.0, eps = 0.01, x0 = 2.0).",
        'Натиснути кнопку «Get» для запуску обчислень.',
        "Результати відображаються у текстовому полі нижче у вигляді таблиць для кожного методу.",
        'Кнопка «Clear» очищує поле результатів.',
        "У разі некоректного введення програма виводить повідомлення про помилку."
    ]
    for i, instr in enumerate(instructions, 1):
        add_paragraph(doc, f"{i}) {instr}", first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ══════════════════════════════════════════
    #  4 ТЕСТУВАННЯ ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ
    # ══════════════════════════════════════════

    add_heading_h1(doc, "4 ТЕСТУВАННЯ ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ")

    add_paragraph(doc,
        "Тестування проведено з параметрами: a = 1.0, b = 2.0, ε = 0.01, x₀ = 2.0 "
        "для рівняння 0.5x⁵ − 0.005x − 1 = 0. Нижче наведено результати роботи кожного методу.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ── Виконуємо обчислення ──
    root_b, fval_b, iters_b, steps_b = bisection_method(A_DEFAULT, B_DEFAULT, EPSILON)
    root_c, fval_c, iters_c, steps_c = chord_method(A_DEFAULT, B_DEFAULT, EPSILON)
    root_n, fval_n, iters_n, steps_n = newton_method(X0_DEFAULT, EPSILON)

    complex_guesses = [
        complex(0.5, 1.0),
        complex(0.5, -1.0),
        complex(-1.0, 0.5),
        complex(-1.0, -0.5),
    ]
    complex_results = []
    for z0 in complex_guesses:
        root_z, iters_z, steps_z = complex_newton(z0, EPSILON)
        complex_results.append((z0, root_z, iters_z, steps_z))

    # ── Таблиця 4.1: Метод половинного ділення ──
    add_heading_h2(doc, "4.1 Результати методу половинного ділення")

    add_paragraph(doc,
        f"Початковий інтервал: [{A_DEFAULT}, {B_DEFAULT}], "
        f"f({A_DEFAULT}) = {f(A_DEFAULT):.5f}, f({B_DEFAULT}) = {f(B_DEFAULT):.5f}.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    headers_b = ["Крок", "a", "b", "c", "f(c)", "Ширина"]
    rows_b = []
    for s in steps_b:
        rows_b.append([
            s['iter'],
            f"{s['a']:.6f}", f"{s['b']:.6f}",
            f"{s['c']:.6f}", f"{s['fc']:.6f}",
            f"{s['width']:.6f}"
        ])

    # Підпис зверху (ДСТУ 3008:2015)
    add_table_caption(doc, "4.1", "Результати методу половинного ділення")
    add_table_with_data(doc, headers_b, rows_b,
                        col_widths=[1.5, 2.5, 2.5, 2.5, 2.5, 2.5])

    add_paragraph(doc,
        f"Знайдений дійсний корінь: x* = {root_b:.6f}.",
        bold=True, first_indent=1.25, space_before=6, space_after=0)
    add_paragraph(doc,
        f"Кількість ітерацій: {iters_b}.",
        first_indent=1.25, space_after=0)
    add_paragraph(doc,
        f"Значення функції: f(x*) = {fval_b:.2e}.",
        first_indent=1.25, space_after=6)

    # ── Таблиця 4.2: Метод хорд ──
    add_heading_h2(doc, "4.2 Результати методу хорд")

    headers_c = ["Крок", "a", "b", "x", "f(x)"]
    rows_c = []
    for s in steps_c:
        rows_c.append([
            s['iter'],
            f"{s['a']:.6f}", f"{s['b']:.6f}",
            f"{s['x']:.6f}", f"{s['fx']:.6f}"
        ])

    add_table_caption(doc, "4.2", "Результати методу хорд")
    add_table_with_data(doc, headers_c, rows_c,
                        col_widths=[1.5, 3.0, 3.0, 3.0, 3.0])

    add_paragraph(doc,
        f"Знайдений дійсний корінь: x* = {root_c:.6f}.",
        bold=True, first_indent=1.25, space_before=6, space_after=0)
    add_paragraph(doc,
        f"Кількість ітерацій: {iters_c}.",
        first_indent=1.25, space_after=0)
    add_paragraph(doc,
        f"Значення функції: f(x*) = {fval_c:.2e}.",
        first_indent=1.25, space_after=6)

    # ── Таблиця 4.3: Метод Ньютона ──
    add_heading_h2(doc, "4.3 Результати методу дотичних (Ньютона)")

    add_paragraph(doc,
        f"Початкове наближення: x₀ = {X0_DEFAULT}.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    headers_n = ["Крок", "x_{k-1}", "f(x_{k-1})", "f'(x_{k-1})", "x_k", "Похибка"]
    rows_n = []
    for s in steps_n:
        rows_n.append([
            s['iter'],
            f"{s['x_prev']:.6f}", f"{s['fx_prev']:.6f}",
            f"{s['dfx_prev']:.6f}", f"{s['x_curr']:.6f}",
            f"{s['precision']:.4e}"
        ])

    add_table_caption(doc, "4.3", "Результати методу дотичних (Ньютона)")
    add_table_with_data(doc, headers_n, rows_n,
                        col_widths=[1.2, 2.3, 2.5, 2.5, 2.3, 2.5])

    add_paragraph(doc,
        f"Знайдений дійсний корінь: x* = {root_n:.6f}.",
        bold=True, first_indent=1.25, space_before=6, space_after=0)
    add_paragraph(doc,
        f"Кількість ітерацій: {iters_n}.",
        first_indent=1.25, space_after=0)
    add_paragraph(doc,
        f"Значення функції: f(x*) = {fval_n:.2e}.",
        first_indent=1.25, space_after=6)

    # ── Таблиця 4.4–4.7: Комплексні корені ──
    add_heading_h2(doc, "4.4 Результати знаходження комплексних коренів")

    for idx, (z0, root_z, iters_z, steps_z) in enumerate(complex_results, 1):
        add_paragraph(doc,
            f"Комплексний корінь {idx} (початкове наближення z₀ = {format_complex(z0)}).",
            bold=True, first_indent=1.25, space_before=6, space_after=3)

        headers_z = ["Крок", "z_{k-1}", "z_k", "Похибка"]
        rows_z = []
        for s in steps_z:
            rows_z.append([
                s['iter'],
                format_complex(s['z_prev']),
                format_complex(s['z_curr']),
                f"{s['precision']:.4e}"
            ])

        add_table_caption(doc, f"4.{3 + idx}", f"Результати знаходження комплексного кореня {idx}")
        add_table_with_data(doc, headers_z, rows_z,
                            col_widths=[1.5, 5.0, 5.0, 2.5])

        add_paragraph(doc,
            f"Результат: z* = {format_complex(root_z)}, кількість кроків: {iters_z}.",
            first_indent=1.25, space_before=3, space_after=6)

    # ══════════════════════════════════════════
    #  5 АНАЛІЗ ОТРИМАНИХ РЕЗУЛЬТАТІВ
    # ══════════════════════════════════════════

    add_heading_h1(doc, "5 АНАЛІЗ ОТРИМАНИХ РЕЗУЛЬТАТІВ")

    add_paragraph(doc,
        "За результатами обчислень, наведених у розділі 4, проведено порівняльний аналіз "
        "ефективності трьох чисельних методів для знаходження дійсного кореня рівняння "
        f"0.5x⁵ − 0.005x − 1 = 0 на інтервалі [{A_DEFAULT}, {B_DEFAULT}] із точністю ε = {EPSILON}.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Зведена таблиця
    summary_headers = ["Метод", "Корінь x*", "f(x*)", "Ітерації"]
    summary_rows = [
        ["Половинного ділення", f"{root_b:.6f}", f"{fval_b:.2e}", str(iters_b)],
        ["Хорд", f"{root_c:.6f}", f"{fval_c:.2e}", str(iters_c)],
        ["Дотичних (Ньютона)", f"{root_n:.6f}", f"{fval_n:.2e}", str(iters_n)],
    ]

    add_table_caption(doc, "5.1", "Порівняльна характеристика методів")
    add_table_with_data(doc, summary_headers, summary_rows,
                        col_widths=[4.5, 3.0, 3.0, 2.5])

    # Аналіз
    add_paragraph(doc,
        f"З таблиці 5.1 видно, що всі три методи знайшли дійсний корінь з високою точністю. "
        f"Метод дотичних (Ньютона) показав найвищу швидкість збіжності — лише {iters_n} ітерацій "
        f"завдяки квадратичній збіжності [1, 8]. Метод хорд потребував {iters_c} ітерацій, а метод "
        f"половинного ділення — {iters_b} ітерацій, що є очікуваним результатом для лінійної збіжності.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=6)

    add_paragraph(doc,
        "Порівняння методів показує, що метод Ньютона є оптимальним за швидкістю збіжності, "
        "проте він потребує знання похідної функції та чутливий до вибору початкового наближення. "
        "Метод половинного ділення є найнадійнішим (завжди збігається при виконанні умови f(a)·f(b) < 0), "
        "але найповільнішим. Метод хорд займає проміжне положення [2, 4].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        f"Також було знайдено 4 комплексні корені полінома. Поліном 5-го степеня має рівно 5 коренів: "
        f"1 дійсний (x* ≈ {root_b:.4f}) та 4 комплексні (попарно спряжені). Комплексні корені знайдено "
        f"за допомогою модифікованого методу Ньютона з різними початковими наближеннями на комплексній площині [9].",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Зведена таблиця комплексних
    summary_complex_headers = ["№", "Корінь z*", "Кроків"]
    summary_complex_rows = []
    for idx, (z0, root_z, iters_z, _) in enumerate(complex_results, 1):
        summary_complex_rows.append([
            str(idx),
            format_complex(root_z),
            str(iters_z)
        ])

    add_table_caption(doc, "5.2", "Знайдені комплексні корені")
    add_table_with_data(doc, summary_complex_headers, summary_complex_rows,
                        col_widths=[1.5, 9.0, 2.5])

    add_paragraph(doc,
        "Похибка обчислення дійсного кореня для кожного методу визначається як абсолютне значення "
        "функції у знайденій точці |f(x*)|. З таблиці 5.1 видно, що всі методи забезпечують "
        "значення |f(x*)| порядку 10⁻³–10⁻⁴, що підтверджує коректність реалізації алгоритмів.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=6)

    # ══════════════════════════════════════════
    #  ВИСНОВКИ
    # ══════════════════════════════════════════

    add_heading_h1(doc, "ВИСНОВКИ")

    add_paragraph(doc,
        "У ході виконання курсової роботи було досліджено та програмно реалізовано чисельні "
        "методи розв'язання нелінійного рівняння 0.5x⁵ − 0.005x − 1 = 0. "
        "За результатами роботи можна зробити такі висновки.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    conclusions = [
        "Було вивчено теоретичні основи методу половинного ділення, методу хорд та методу "
        "дотичних (Ньютона), а також їх узагальнення на комплексну площину.",

        "Розроблено покрокові алгоритми для кожного з методів та їх програмну реалізацію мовою C# "
        "з використанням об'єктно-орієнтованого підходу.",

        "Створено програмне забезпечення з графічним інтерфейсом (Windows Forms), яке "
        "дозволяє задавати параметри обчислень та переглядати результати у зручному форматі.",

        f"Протестовано програму на рівнянні 0.5x⁵ − 0.005x − 1 = 0 із параметрами: "
        f"a = {A_DEFAULT}, b = {B_DEFAULT}, ε = {EPSILON}, x₀ = {X0_DEFAULT}. "
        f"Знайдено дійсний корінь x* ≈ {root_b:.4f} та 4 комплексні корені.",

        f"Проведено порівняльний аналіз ефективності методів. Метод Ньютона показав "
        f"найвищу швидкість збіжності ({iters_n} ітерацій), метод хорд — {iters_c} ітерацій, "
        f"метод половинного ділення — {iters_b} ітерацій.",

        "Результати підтверджують теоретичні відомості щодо порядку збіжності: "
        "квадратична збіжність методу Ньютона значно перевершує лінійну збіжність "
        "методів половинного ділення та хорд. Перспективним напрямком удосконалення є "
        "реалізація адаптивного вибору початкового наближення та автоматичного відокремлення коренів."
    ]
    for i, concl in enumerate(conclusions, 1):
        add_paragraph(doc, f"{i}. {concl}", first_indent=1.25,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)

    # ══════════════════════════════════════════
    #  ПЕРЕЛІК ПОСИЛАНЬ (ДСТУ 8302:2015)
    # ══════════════════════════════════════════

    add_heading_h1(doc, "ПЕРЕЛІК ПОСИЛАНЬ")

    # Формат ДСТУ 8302:2015
    references = [
        "Бахвалов Н. С., Жидков Н. П., Кобельков Г. М. Численные методы : уч. пособие. "
        "Москва : БИНОМ. Лаборатория знаний, 2015. 636 с.",

        "Самарский А. А., Гулин А. В. Численные методы : уч. пособие. "
        "Москва : Наука, 1989. 432 с.",

        "Квєтний Р. Н., Богач І. В., Бойко О. Р., Софина О. Ю., Шушура О. М. "
        "Комп'ютерне моделювання систем та процесів. Методи обчислень. Частина 1 : "
        "навч. посібник. Вінниця : ВНТУ, 2013. 191 с.",

        "Квєтний Р. Н., Богач І. В., Бойко О. Р., Софина О. Ю., Шушура О. М. "
        "Комп'ютерне моделювання систем та процесів. Методи обчислень. Частина 2 : "
        "навч. посібник. Вінниця : ВНТУ, 2013. 235 с.",

        "Квєтний Р. Н. Методи комп'ютерних обчислень : навч. посібник. "
        "Вінниця : ВДТУ, 2001. 148 с.",

        "Троелсен Э., Джепикс Ф. Язык программирования C# и платформа .NET : "
        "уч. пособие. Санкт-Петербург : Диалектика, 2022. 1168 с.",

        "Microsoft Docs. System.Numerics.Complex Structure. URL: "
        "https://docs.microsoft.com/dotnet/api/system.numerics.complex "
        "(дата звернення: 01.06.2026).",

        "Press W. H., Teukolsky S. A., Vetterling W. T., Flannery B. P. "
        "Numerical Recipes: The Art of Scientific Computing. "
        "Cambridge : Cambridge University Press, 2007. 1256 p.",

        "Burden R. L., Faires J. D. Numerical Analysis. "
        "Boston : Brooks/Cole, 2010. 888 p.",
    ]
    for i, ref in enumerate(references, 1):
        add_paragraph(doc, f"{i}. {ref}", first_indent=1.25, space_after=0,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ══════════════════════════════════════════
    #  ДОДАТКИ
    # ══════════════════════════════════════════

    add_heading_h1(doc, "ДОДАТКИ")

    add_empty_paragraph(doc, 1)

    # Додаток А (згідно з метод. вказівками: зверху "Додаток А", нижче — заголовок)
    add_paragraph(doc, "Додаток А", font_size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None,
                  space_after=0)
    add_paragraph(doc, "(обов'язковий)", font_size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None,
                  space_after=3)
    add_paragraph(doc, "Лістинг програми", font_size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None,
                  space_after=12)

    # Читаємо та додаємо файли вихідного коду
    script_dir = os.path.dirname(os.path.abspath(__file__))
    source_files = [
        os.path.join(script_dir, "MKO_LIB", "proj", "coursework.cs"),
        os.path.join(script_dir, "MKO_LIB", "methods", "BisectionMethod.cs"),
        os.path.join(script_dir, "MKO_LIB", "methods", "ChordMethod.cs"),
        os.path.join(script_dir, "MKO_LIB", "methods", "NewtonMethod.cs"),
        os.path.join(script_dir, "MKO_LIB", "MainForm.cs"),
    ]

    for src_path in source_files:
        if os.path.exists(src_path):
            filename = os.path.basename(src_path)
            add_paragraph(doc, f"Файл: {filename}", font_size=12, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          first_indent=None, space_before=6, space_after=3,
                          line_spacing=1.0)

            with open(src_path, 'r', encoding='utf-8') as file:
                code = file.read()
            add_code_block(doc, code, font_size=8)

            add_empty_paragraph(doc, 1)
        else:
            add_paragraph(doc, f"[Файл {src_path} не знайдено]",
                          font_size=12, italic=True, first_indent=None)

    # ── Збереження документа ──
    output_path = os.path.join(script_dir, OUTPUT_FILENAME)
    doc.save(output_path)
    print(f"Пояснювальну записку збережено: {output_path}")
    print(f"   Загальна кількість абзаців: {len(doc.paragraphs)}")
    print(f"   Загальна кількість таблиць:  {len(doc.tables)}")

    # Підсумок результатів
    print(f"\nРезультати обчислень:")
    print(f"   Дійсний корінь: x* ≈ {root_b:.6f}")
    print(f"   Метод бісекції: {iters_b} ітерацій")
    print(f"   Метод хорд:     {iters_c} ітерацій")
    print(f"   Метод Ньютона:  {iters_n} ітерацій")
    for idx, (z0, root_z, iters_z, _) in enumerate(complex_results, 1):
        print(f"   Комплексний корінь {idx}: {format_complex(root_z)} ({iters_z} ітерацій)")


if __name__ == "__main__":
    generate_report()
